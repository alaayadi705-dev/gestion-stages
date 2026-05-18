from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import docx.oxml.ns as ns
from docx.oxml.ns import qn

def create_memoire():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ==========================================
    # TITLE PAGE
    # ==========================================
    for _ in range(4):
        doc.add_paragraph('')
    
    # Header info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('République Tunisienne')
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Ministère de l\'Enseignement Supérieur\net de la Recherche Scientifique')
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Université de Kairouan')
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Faculté des Sciences et Techniques\nde Sidi Bouzid')
    run.bold = True
    run.font.size = Pt(12)
    
    for _ in range(2):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('N° d\'ordre : ………')
    run.bold = True
    run.font.size = Pt(14)
    
    for _ in range(3):
        doc.add_paragraph('')
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MÉMOIRE DE PROJET DE FIN D\'ÉTUDES')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('En vue de l\'obtention du\nDiplôme de License en Informatique')
    run.font.size = Pt(14)
    
    for _ in range(2):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Présenté par')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Ayadi Ala Eddine')
    run.bold = True
    run.font.size = Pt(20)
    
    for _ in range(2):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Sujet :')
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Guide de conduite d\'un projet informatique\nselon le cadre de développement SCRUM')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Application : Système de Gestion des Stages à l\'Étranger')
    run.font.size = Pt(13)
    
    for _ in range(3):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Soutenu en Juin 2026, devant le jury composé de :')
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('M./Mme ……………………………… Président(e)\nM./Mme ……………………………… Rapporteur\nM./Mme ……………………………… Encadrant(e)')
    run.font.size = Pt(12)
    
    for _ in range(2):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A.U. : 2025-2026')
    run.bold = True
    run.font.size = Pt(14)
    
    for _ in range(3):
        doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Lieu de stage : Centre National de l\'Informatique (CNI)')
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ==========================================
    # DÉDICACES
    # ==========================================
    doc.add_heading('DÉDICACES', level=1)
    doc.add_paragraph('')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('À mes chers parents,\n').italic = True
    p.add_run('Pour leurs sacrifices, leur amour et leur soutien inconditionnel.\n\n').italic = True
    p.add_run('À ma famille,\n').italic = True
    p.add_run('Pour leurs encouragements constants.\n\n').italic = True
    p.add_run('À mes amis et collègues,\n').italic = True
    p.add_run('Pour leur accompagnement tout au long de ce parcours.\n\n').italic = True
    p.add_run('À tous ceux qui croient en la puissance de la technologie\n').italic = True
    p.add_run('pour transformer le monde.').italic = True
    
    doc.add_page_break()
    
    # ==========================================
    # REMERCIEMENTS
    # ==========================================
    doc.add_heading('REMERCIEMENTS', level=1)
    doc.add_paragraph('')
    
    doc.add_paragraph(
        'Je tiens tout d\'abord à exprimer ma profonde gratitude à '
        'Dieu le tout puissant de m\'avoir donné la force, la patience '
        'et la volonté pour mener à bien ce projet.'
    )
    
    doc.add_paragraph(
        'Je souhaite adresser mes plus sincères remerciements à '
        'l\'ensemble du corps enseignant de la Faculté des Sciences '
        'et Techniques de Sidi Bouzid, pour la qualité de la formation '
        'qu\'ils m\'ont dispensée tout au long de mon cursus universitaire.'
    )
    
    doc.add_paragraph(
        'Mes remerciements vont particulièrement à '
        'mon encadrant académique, pour ses précieux conseils, '
        'sa disponibilité et son accompagnement tout au long de '
        'la réalisation de ce projet.'
    )
    
    doc.add_paragraph(
        'Je tiens également à remercier le Centre National de l\'Informatique (CNI) '
        'pour m\'avoir accueilli en stage et offert un environnement de travail '
        'propice à l\'apprentissage et à l\'innovation.'
    )
    
    doc.add_paragraph(
        'Je remercie chaleureusement tous les professionnels du CNI '
        'qui m\'ont fait part de leur expertise et m\'ont guidé dans la '
        'compréhension des enjeux métier liés à la gestion des stages à l\'étranger.'
    )
    
    doc.add_paragraph(
        'Enfin, je remercie ma famille et mes amis pour leur soutien moral '
        'et leurs encouragements constants durant cette période de formation '
        'et de réalisation du projet.'
    )
    
    doc.add_page_break()
    
    # ==========================================
    # TABLE DES MATIÈRES (placeholder)
    # ==========================================
    doc.add_heading('TABLE DES MATIÈRES', level=1)
    doc.add_paragraph('')
    doc.add_paragraph('(La table des matières sera générée automatiquement par Word)')
    doc.add_paragraph('')
    doc.add_paragraph('Allez dans Références > Table des matières > Table automatique')
    
    doc.add_page_break()
    
    # ==========================================
    # INTRODUCTION GÉNÉRALE
    # ==========================================
    doc.add_heading('INTRODUCTION GÉNÉRALE', level=1)
    doc.add_paragraph('')
    
    doc.add_paragraph(
        'La mobilité internationale des stagiaires représente aujourd\'hui un enjeu majeur '
        'pour les institutions académiques et gouvernementales tunisiennes. Dans un contexte '
        'de mondialisation croissante, les opportunités de stages à l\'étranger offrent aux '
        'étudiants et professionnels une expérience invaluable pour développer leurs compétences, '
        'enrichir leur culture et renforcer leur employabilité sur le marché international.'
    )
    
    doc.add_paragraph(
        'Cependant, la gestion administrative de ces stages internationaux demeure un processus '
        'complexe, impliquant de multiples acteurs : les stagiaires eux-mêmes, les établissements '
        'd\'enseignement, les entreprises d\'accueil à l\'étranger, les ministères de tutelle, et '
        'les organismes de financement. Chaque étape du processus – de la candidature à la validation '
        'du rapport de stage en passant par la gestion des frais et le suivi administratif – génère '
        'une documentation importante qui doit être organisée, tracée et archivée de manière rigoureuse.'
    )
    
    doc.add_paragraph(
        'Avant la mise en place de notre solution, le Centre National de l\'Informatique (CNI) '
        'gérait ces processus principalement à travers des outils bureautiques classiques (tableurs, '
        'documents Word) et des échanges par email. Cette approche présentait plusieurs limitations : '
        'dispersion de l\'information, difficulté de suivi en temps réel, risques d\'erreurs lors de '
        'la saisie manuelle, absence de statistiques consolidées, et manque de traçabilité des validations.'
    )
    
    doc.add_paragraph(
        'Face à ces constats, il est apparu nécessaire de concevoir et développer une application web '
        'dédiée à la gestion des stages à l\'étranger, permettant de centraliser l\'ensemble des '
        'informations, d\'automatiser les processus métier, et de fournir des outils de pilotage et '
        'de décision aux différents acteurs impliqués.'
    )
    
    doc.add_paragraph(
        'Ce projet s\'inscrit dans le cadre de l\'obtention du Diplôme de License en Informatique à '
        'la Faculté des Sciences et Techniques de Sidi Bouzid. Il a été réalisé au sein du Centre '
        'National de l\'Informatique (CNI) sur une durée de quatre mois, en suivant le framework de '
        'gestion de projets agiles SCRUM.'
    )
    
    doc.add_paragraph(
        'L\'objectif principal de ce mémoire est de documenter l\'ensemble du processus de développement '
        'de cette application, depuis l\'analyse des besoins jusqu\'à la réalisation finale, en passant '
        'par les choix méthodologiques, architecturaux et techniques.'
    )
    
    doc.add_paragraph('Ce document est structuré en quatre chapitres principaux :')
    
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    p.add_run('Le Chapitre 1 ').bold = True
    p.add_run('présente le cadre général du projet, incluant l\'organisme d\'accueil, l\'étude de l\'existant, la critique de celui-ci, et les choix méthodologiques adoptés, notamment le framework SCRUM et la méthodologie orientée objet.')
    
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    p.add_run('Le Chapitre 2 ').bold = True
    p.add_run('détaille l\'analyse préliminaire du projet, avec la présentation des rôles SCRUM, le Product Backlog initial, l\'architecture technique de l\'application, et l\'environnement de travail mis en place.')
    
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    p.add_run('Le Chapitre 3 ').bold = True
    p.add_run('décrit le déroulé des sprints de développement, incluant la planification globale, un exemple détaillé d\'un sprint représentatif, et la gestion des obstacles rencontrés.')
    
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    p.add_run('Le Chapitre 4 ').bold = True
    p.add_run('présente les résultats finaux obtenus, les livrables produits, les démonstrations des fonctionnalités implémentées, ainsi que les tests réalisés et les améliorations apportées.')
    
    doc.add_paragraph(
        'Enfin, nous conclurons ce mémoire par une synthèse générale présentant les apports du projet, '
        'les difficultés rencontrées, et les perspectives d\'évolution de l\'application.'
    )
    
    # Save document
    output_path = r'c:\Users\alaay\OneDrive\Bureau\projet pfe\gestion_projets\gestionstages\memoire_pfe.docx'
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")
    print("Part 1 complete. Continue with Part 2 for chapters...")

if __name__ == '__main__':
    create_memoire()
